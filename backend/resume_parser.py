"""Resume extraction: text, contact details and embedded hyperlinks.

Supports PDF, DOCX, plain text, and (best-effort) images. The design goal is to
fix the reported email mis-parse at its root and to pull the links that resumes
usually embed rather than print.

Why the email bug happened: PDF text extraction collapses the whitespace
between adjacent text runs, so a label like "Skype"/"Profile" glues onto the
address and the greedy email regex captures the trailing letters of the
previous word (e.g. dksmayank03@… → pedksmayank03@…). The fix is to prefer the
exact address from the PDF's embedded `mailto:` link annotation, which carries
the real value regardless of how the visible text was laid out. The same
annotations give us the LinkedIn/GitHub/portfolio URLs.

All heavy or optional dependencies (python-docx, OCR) are imported lazily inside
the functions that use them, so this module imports cleanly even when they are
not installed — an unsupported file simply yields no text rather than crashing.
"""

import io
import logging
import os
import re

logger = logging.getLogger(__name__)

# Local part kept conservative; the annotation path is the authoritative source.
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\s().\-]{7,}\d)")
URL_RE = re.compile(r"https?://[^\s)>\]\"']+", re.IGNORECASE)

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".webp")
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp")


# ---------------------------------------------------------------------------
# Per-format extractors — each returns (text, [embedded_link_uris])
# ---------------------------------------------------------------------------

# Raised when a file is recognised but its text genuinely cannot be extracted
# (encrypted/corrupt), so the caller can tell the user rather than store a blank.
class ExtractionError(Exception):
    pass


def _extract_pdf(content: bytes) -> tuple[str, list]:
    text_parts, links = [], []
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        logger.warning("Could not open PDF: %s", exc)
        raise ExtractionError("This PDF couldn't be opened — it may be corrupt.") from exc

    # Password-protected PDFs: try an empty owner password, else report clearly.
    if getattr(reader, "is_encrypted", False):
        try:
            if reader.decrypt("") == 0:
                raise ExtractionError("This PDF is password-protected. Remove the password and re-upload.")
        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError("This PDF is password-protected. Remove the password and re-upload.") from exc

    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            pass
        # Link annotations hold the exact URI (mailto:, https://…) that the
        # visible text may have mangled.
        try:
            annots = page.get("/Annots")
            if not annots:
                continue
            for ref in annots:
                obj = ref.get_object()
                action = obj.get("/A")
                if action:
                    uri = action.get("/URI")
                    if uri:
                        links.append(str(uri))
        except Exception:
            pass

    return "\n".join(text_parts).strip(), links


def _extract_docx(content: bytes) -> tuple[str, list]:
    try:
        from docx import Document
    except Exception:
        logger.warning("python-docx is not installed; cannot read .docx")
        return "", []

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("Could not open DOCX: %s", exc)
        return "", []

    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    links = []
    try:
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype and rel._target:
                links.append(str(rel._target))
    except Exception:
        pass

    return "\n".join(parts).strip(), links


def _extract_doc_binary(content: bytes) -> tuple[str, list]:
    """Best-effort text from a legacy binary .doc (Word 97-2003).

    There is no lightweight pure-Python reader for the old .doc format, so we
    pull readable text runs directly from the bytes. A .doc stores its text as
    either UTF-16LE or CP-1252 runs interleaved with binary structures; we decode
    both ways and keep whichever yields more real, alphabetic lines. Messy but
    non-empty — the AI structuring step and the recruiter's edits do the rest,
    which is far better than silently storing a blank."""
    def readable(decoded: str) -> str:
        cleaned = "".join(c if (c.isprintable() or c in "\n\t ") else "\n" for c in decoded)
        lines = []
        for ln in cleaned.split("\n"):
            ln = ln.strip()
            # Keep lines that look like words, not stray formatting tokens.
            if len(ln) >= 4 and sum(c.isalpha() for c in ln) >= max(3, len(ln) // 2):
                lines.append(ln)
        return "\n".join(lines).strip()

    best = ""
    for enc in ("utf-16-le", "cp1252", "latin-1"):
        try:
            candidate = readable(content.decode(enc, errors="ignore"))
        except Exception:
            candidate = ""
        if len(candidate) > len(best):
            best = candidate
    return best, []


def _extract_text_file(content: bytes) -> tuple[str, list]:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(enc).strip(), []
        except Exception:
            continue
    return "", []


def _ocr_available() -> bool:
    """True only if a free OCR stack is importable AND usable on this host."""
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except Exception:
        return False
    try:
        # pytesseract needs the Tesseract binary; this raises if it's absent.
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _extract_image(content: bytes) -> tuple[str, list]:
    """Best-effort OCR. Returns empty text (not an error) when OCR is unavailable,
    so image uploads still create a candidate record that can be edited by hand."""
    if not _ocr_available():
        logger.info("OCR not available on this host; storing image resume without text")
        return "", []
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        return (pytesseract.image_to_string(img) or "").strip(), []
    except Exception as exc:
        logger.warning("OCR failed: %s", exc)
        return "", []


# ---------------------------------------------------------------------------
# Link classification
# ---------------------------------------------------------------------------

def _classify_links(links: list) -> dict:
    """Split raw URIs into LinkedIn / GitHub / portfolio and a clean list."""
    seen, cleaned = set(), []
    linkedin = github = portfolio = None

    for raw in links:
        if not raw:
            continue
        url = raw.strip()
        low = url.lower()
        if low.startswith("mailto:") or low.startswith("tel:"):
            continue
        if url in seen:
            continue
        seen.add(url)
        cleaned.append(url)

        if not linkedin and "linkedin.com" in low:
            linkedin = url
        elif not github and "github.com" in low:
            github = url
        elif not portfolio and low.startswith("http"):
            # First non-social link is treated as a portfolio/website.
            portfolio = url

    return {"links": cleaned, "linkedin": linkedin, "github": github, "portfolio": portfolio}


def _best_email(text: str, links: list) -> str | None:
    # 1) Exact address from a mailto: annotation — immune to text mangling.
    for raw in links:
        low = (raw or "").lower()
        if low.startswith("mailto:"):
            match = EMAIL_RE.search(raw)
            if match:
                return match.group(0).lower()
    # 2) Fall back to scanning the extracted text.
    match = EMAIL_RE.search(text or "")
    return match.group(0).lower() if match else None


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_document(content: bytes, filename: str) -> dict:
    """Extract everything we can from one resume file.

    Returns: text, email, phone, links (list), linkedin, github, portfolio,
    and `ocr_used`/`unsupported` flags for the caller to surface if needed.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    warning = None

    try:
        if ext == ".pdf":
            text, raw_links = _extract_pdf(content)
        elif ext == ".docx":
            text, raw_links = _extract_docx(content)
        elif ext == ".doc":
            # Some .doc files are actually zipped .docx mislabelled — try that
            # first, then fall back to best-effort binary extraction.
            text, raw_links = _extract_docx(content)
            if not text:
                text, raw_links = _extract_doc_binary(content)
        elif ext == ".txt":
            text, raw_links = _extract_text_file(content)
        elif ext in IMAGE_EXTENSIONS:
            text, raw_links = _extract_image(content)
            if not text:
                warning = "No text could be read from this image. If OCR isn't available, type the details in manually."
        else:
            text, raw_links = "", []
    except ExtractionError as exc:
        text, raw_links, warning = "", [], str(exc)

    # A text-bearing document that yielded nothing is worth flagging so the user
    # can re-upload (e.g. a PDF that is just a scanned image, or an odd .doc).
    if warning is None and ext in (".pdf", ".doc", ".docx", ".txt") and not (text or "").strip():
        warning = "We couldn't read any text from this file. Try re-uploading it as a PDF."

    # Any URLs printed as plain text count too.
    for match in URL_RE.finditer(text or ""):
        raw_links.append(match.group(0))

    classified = _classify_links(raw_links)
    email = _best_email(text, raw_links)
    phone_match = PHONE_RE.search(text or "")
    phone = phone_match.group(0).strip() if phone_match else None

    return {
        "text": text or "",
        "email": email,
        "phone": phone,
        "links": classified["links"],
        "linkedin": classified["linkedin"],
        "github": classified["github"],
        "portfolio": classified["portfolio"],
        # None when text was extracted; a human message when it wasn't, so the
        # upload endpoint can tell the user which files need attention.
        "warning": warning,
        "has_text": bool((text or "").strip()),
    }
